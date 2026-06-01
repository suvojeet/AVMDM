"""
Rebuild USPTO_PATENT_APPLICATION.docx with embedded patent figures.
Run: python build_docx_with_figures.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE = Path(__file__).parent
MD_FILE = BASE / "USPTO_PATENT_APPLICATION.md"
FIGURES_DIR = BASE / "figures"
OUT_FILE = BASE / "USPTO_PATENT_APPLICATION.docx"

# ── Colours ──────────────────────────────────────────────────────────────────
C_INDIGO  = RGBColor(0x3B, 0x00, 0x7F)   # deep indigo — headings
C_VIOLET  = RGBColor(0x60, 0x00, 0xBF)   # violet — sub-headings
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_LGREY   = RGBColor(0xF3, 0xF4, 0xF6)
C_DKBLUE  = RGBColor(0x1E, 0x1B, 0x4B)

FIGURE_FILES = {
    1: FIGURES_DIR / "FIG1_architecture.png",
    2: FIGURES_DIR / "FIG2_blocking.png",
    3: FIGURES_DIR / "FIG3_pipeline.png",
    4: FIGURES_DIR / "FIG4_em_algorithm.png",
    5: FIGURES_DIR / "FIG5_drift_detection.png",
    6: FIGURES_DIR / "FIG6_data_model.png",
    7: FIGURES_DIR / "FIG7_provisional_identity.png",
    8: FIGURES_DIR / "FIG8_ml_feedback.png",
}

FIGURE_CAPTIONS = {
    1: "FIG. 1 — System Architecture: Overall Averio MDM platform showing the blocking engine, three-stage match pipeline, survivorship engine, and dual data stores (Neo4j + Azure Cosmos DB).",
    2: "FIG. 2 — Nine-Strategy Union Blocking: Flowchart of the nine independent blocking key generation strategies applied in union for O(N×k) candidate retrieval.",
    3: "FIG. 3 — Three-Stage Cascading Match Pipeline: Threshold-driven routing through deterministic, probabilistic Fellegi-Sunter, and AI-enhanced LLM matching stages.",
    4: "FIG. 4 — Self-Calibrating Fellegi-Sunter EM Algorithm: E-step and M-step iterations for autonomous m-probability and u-probability estimation across ten attributes.",
    5: "FIG. 5 — Post-Update Cluster Drift Detection: Probabilistic re-evaluation of cluster membership following attribute update, with automated reassignment and steward escalation.",
    6: "FIG. 6 — Hybrid Graph + Document Data Model: Relationships between Party nodes (Neo4j), GoldenRecord documents, and TimelineEvent documents (Azure Cosmos DB).",
    7: "FIG. 7 — Provisional Golden Identity Lifecycle: Assignment of a globally unique golden identifier at first persistence, eliminating the null golden state throughout steward review.",
    8: "FIG. 8 — Human-in-the-Loop ML Feedback Pipeline: Feature vector capture at steward decision time, immutable feedback storage, and automatic logistic regression retraining.",
}


# ── XML helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def page_break(doc: Document):
    doc.add_page_break()


def add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "6B7280")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Document setup ────────────────────────────────────────────────────────────

def setup_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.25)
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Heading styles
    for h, size, color in [
        ("Heading 1", 16, C_INDIGO),
        ("Heading 2", 13, C_VIOLET),
        ("Heading 3", 12, C_DKBLUE),
    ]:
        s = doc.styles[h]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after  = Pt(6)
        s.paragraph_format.keep_with_next = True

    return doc


# ── Cover banner ──────────────────────────────────────────────────────────────

def add_cover(doc: Document):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "3B007F")
    cell.width = Inches(6)

    lines = [
        ("UNITED STATES PATENT APPLICATION", 18, True),
        ("Non-Provisional Utility Patent — 35 U.S.C. § 111(a)", 11, False),
        ("", 6, False),
        ("Applicant: Averio Technologies Inc.", 11, False),
        ("Inventors: Suvojeet Pal; Rakhi Chatterjee", 11, False),
        ("Attorney Docket No.: AVERIO-001-US", 11, False),
    ]
    for text, size, bold in lines:
        p = cell.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.name  = "Arial"
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.color.rgb = C_WHITE
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

    doc.add_paragraph()  # spacing after banner


# ── Paragraph writer ──────────────────────────────────────────────────────────

def apply_inline_bold(para, text: str):
    """Parse **bold** markers and add runs with/without bold."""
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        if i % 2 == 1:   # inside **...**
            run.bold = True


def add_body_para(doc: Document, text: str, indent: float = 0,
                  bold: bool = False, italic: bool = False,
                  space_before: float = 2, space_after: float = 4,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)

    if "**" in text:
        apply_inline_bold(p, text)
    else:
        run = p.add_run(text)
        run.font.name   = "Times New Roman"
        run.font.size   = Pt(12)
        run.bold        = bold
        run.italic      = italic

    return p


def add_code_block(doc: Document, text: str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "F3F4F6")
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    doc.add_paragraph()


# ── Figure insertion ──────────────────────────────────────────────────────────

def insert_figure(doc: Document, fig_num: int):
    """Insert figure image centred with caption below."""
    fig_path = FIGURE_FILES.get(fig_num)
    if not fig_path or not fig_path.exists():
        add_body_para(doc, f"[Figure {fig_num} — file not found: {fig_path}]",
                      italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        return

    # Image paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run()
    run.add_picture(str(fig_path), width=Inches(5.5))

    # Caption paragraph
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after  = Pt(16)
    r = cap.add_run(FIGURE_CAPTIONS[fig_num])
    r.font.name   = "Arial"
    r.font.size   = Pt(9)
    r.font.italic = True
    r.font.color.rgb = C_INDIGO


# ── Drawings section ──────────────────────────────────────────────────────────

def add_drawings_section(doc: Document):
    page_break(doc)
    h = doc.add_heading("DRAWINGS", level=1)
    h.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph("(Referenced in Brief Description of the Drawings)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.runs[0]
    r.font.name   = "Arial"
    r.font.size   = Pt(10)
    r.font.italic = True
    r.font.color.rgb = C_VIOLET
    sub.paragraph_format.space_after = Pt(12)

    for i in range(1, 9):
        add_body_para(doc, f"FIG. {i}", bold=True, space_before=6, space_after=2,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        insert_figure(doc, i)
        if i < 8:
            add_horizontal_rule(doc)
            doc.add_paragraph()


# ── Markdown parser / renderer ────────────────────────────────────────────────

def render_markdown(doc: Document, md_text: str):
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf = []
    in_claims = False

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_buf))
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:].strip()
            # Skip the very first H1 (already in cover banner)
            if text == "UNITED STATES PATENT APPLICATION":
                i += 1
                continue
            doc.add_heading(text, level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            text = stripped[4:].strip()
            doc.add_heading(text, level=3)
            in_claims = "CLAIM" in text.upper()
            i += 1
            continue

        if stripped.startswith("#### "):
            text = stripped[5:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name  = "Arial"
            run.font.size  = Pt(11)
            run.font.bold  = True
            run.font.color.rgb = C_DKBLUE
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # Bold-only line (e.g. **Application Type:** ...)
        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            text = stripped[2:-2]
            add_body_para(doc, text, bold=True, space_before=2, space_after=2)
            i += 1
            continue

        # Metadata key-value lines ("**Key:** Value")
        if stripped.startswith("**") and ":**" in stripped[:40]:
            add_body_para(doc, stripped, space_before=2, space_after=2)
            i += 1
            continue

        # Unordered list
        if re.match(r"^[-*+]\s", stripped):
            text = re.sub(r"^[-*+]\s+", "", stripped)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # strip bold markers for list
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            i += 1
            continue

        # Figure reference lines in Brief Description (e.g. "**FIG. 1** is a system...")
        fig_match = re.match(r"\*\*FIG\. (\d+)\*\*\s+(.*)", stripped)
        if fig_match:
            fig_num   = int(fig_match.group(1))
            remainder = fig_match.group(2)
            full_text = f"FIG. {fig_num} {remainder}"
            add_body_para(doc, full_text, bold=False, space_before=4, space_after=2)
            i += 1
            continue

        # Regular paragraph
        add_body_para(doc, stripped, space_before=2, space_after=4)
        i += 1


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print("Building USPTO patent DOCX with embedded figures...")

    md_text = MD_FILE.read_text(encoding="utf-8")

    doc = setup_document()
    add_cover(doc)

    # Split markdown at "BRIEF DESCRIPTION OF THE DRAWINGS" so we can insert
    # the Drawings section immediately after that section ends.
    DRAWINGS_HEADING = "## BRIEF DESCRIPTION OF THE DRAWINGS"
    DETAILED_HEADING = "## DETAILED DESCRIPTION"

    before_drawings = md_text
    drawings_section = ""
    after_drawings = ""

    idx_d = md_text.find(DRAWINGS_HEADING)
    idx_det = md_text.find(DETAILED_HEADING)

    if idx_d != -1 and idx_det != -1:
        before_drawings   = md_text[:idx_d]
        drawings_section  = md_text[idx_d:idx_det]
        after_drawings     = md_text[idx_det:]
    elif idx_d != -1:
        before_drawings  = md_text[:idx_d]
        drawings_section = md_text[idx_d:]
        after_drawings   = ""

    # Render everything before "Brief Description"
    render_markdown(doc, before_drawings)

    # Render "Brief Description of the Drawings" text
    render_markdown(doc, drawings_section)

    # ── Insert the full Drawings section (all 8 figures) ──────────────────────
    add_drawings_section(doc)

    # Render the rest of the patent (Detailed Description, Claims, Abstract, etc.)
    if after_drawings:
        page_break(doc)
        render_markdown(doc, after_drawings)

    doc.save(str(OUT_FILE))
    print(f"Saved: {OUT_FILE}")
    print(f"File size: {OUT_FILE.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
